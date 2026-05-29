import logging
import re
import fitz  # pymupdf
from pathlib import Path

logger = logging.getLogger(__name__)

# Regex for detecting section headers: short line, title-case or ALL CAPS, no trailing period
_HEADER_RE = re.compile(r"^(?:[A-Z][A-Za-z0-9 ,\-]{0,60}[^.]|[A-Z]{2,}[A-Z0-9 ,\-]{0,60})$")


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(pdf_path: str) -> list[tuple[int, str]]:
    """Return list of (1-based page_number, page_text) with tables rendered as markdown."""
    doc = fitz.open(pdf_path)
    pages: list[tuple[int, str]] = []
    for i, page in enumerate(doc, 1):
        text = page.get_text()
        # Append any tables detected on the page as markdown blocks
        try:
            for table in page.find_tables().tables:
                rows = table.extract()
                if rows:
                    md_rows = [
                        "| " + " | ".join(str(cell or "").strip() for cell in row) + " |"
                        for row in rows
                    ]
                    text += "\n\n" + "\n".join(md_rows)
        except Exception:
            pass  # table extraction is best-effort
        pages.append((i, text))
    doc.close()
    total_chars = sum(len(t) for _, t in pages)
    logger.info("Extracted %d chars across %d pages from PDF", total_chars, len(pages))
    return pages


def extract_text_from_docx(docx_path: str) -> list[tuple[int, str]]:
    """Extract text from a Word document. Returns single page (no page info in docx)."""
    from docx import Document as DocxDocument
    doc = DocxDocument(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append("| " + " | ".join(cells) + " |")
    text = "\n\n".join(paragraphs)
    logger.info("Extracted %d chars from DOCX", len(text))
    return [(1, text)]


def extract_text_from_txt(txt_path: str) -> list[tuple[int, str]]:
    """Extract text from a plain-text file."""
    text = Path(txt_path).read_text(encoding="utf-8", errors="replace")
    logger.info("Extracted %d chars from TXT", len(text))
    return [(1, text)]


def extract_text(file_path: str) -> list[tuple[int, str]]:
    """Dispatch extraction based on file extension."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    if ext == ".txt":
        return extract_text_from_txt(file_path)
    raise ValueError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# Semantic chunking
# ---------------------------------------------------------------------------

def _is_header(line: str) -> bool:
    line = line.strip()
    if len(line) < 3 or len(line) > 80:
        return False
    return bool(_HEADER_RE.match(line))


def split_into_chunks(
    pages: list[tuple[int, str]],
    chunk_size: int,
    overlap: int,
) -> list[dict]:
    """
    Semantic-aware chunking that:
    - Treats section headers as hard split points
    - Prefers paragraph (double-newline) boundaries
    - Falls back to sentence then word boundaries
    Returns list of {"content": str, "page_number": int}.
    """
    # Build full text preserving page boundaries
    full_text = ""
    page_boundaries: list[tuple[int, int]] = []  # (start_char, page_number)

    for page_num, page_text in pages:
        page_boundaries.append((len(full_text), page_num))
        full_text += page_text + "\n\n"

    full_text = re.sub(r"\n{3,}", "\n\n", full_text).strip()

    def get_page(pos: int) -> int:
        page = page_boundaries[0][1] if page_boundaries else 1
        for start, pnum in page_boundaries:
            if start <= pos:
                page = pnum
            else:
                break
        return page

    # Identify header positions as hard split points
    hard_breaks: set[int] = set()
    for m in re.finditer(r"(?:^|\n)([^\n]{3,80})\n", full_text):
        line = m.group(1).strip()
        if _is_header(line):
            hard_breaks.add(m.start())

    chunks: list[dict] = []
    start = 0

    while start < len(full_text):
        end = start + chunk_size

        if end >= len(full_text):
            chunk_text = full_text[start:].strip()
            if len(chunk_text) > 30:
                chunks.append({"content": chunk_text, "page_number": get_page(start)})
            break

        # Check for a hard break (header) within the window
        boundary = -1
        for hb in sorted(hard_breaks):
            if start < hb <= end:
                boundary = hb
                break

        if boundary == -1:
            # Prefer paragraph > sentence > word boundaries
            for sep in ["\n\n", ". ", "? ", "! ", "\n", " "]:
                pos = full_text.rfind(sep, start, end)
                if pos != -1:
                    boundary = pos + len(sep)
                    break

        if boundary == -1 or boundary <= start:
            boundary = end

        chunk_text = full_text[start:boundary].strip()
        if len(chunk_text) > 30:
            chunks.append({"content": chunk_text, "page_number": get_page(start)})

        next_start = boundary - overlap
        start = next_start if next_start > start else boundary

    logger.info("Split into %d semantic chunks (size=%d, overlap=%d)", len(chunks), chunk_size, overlap)
    return chunks
